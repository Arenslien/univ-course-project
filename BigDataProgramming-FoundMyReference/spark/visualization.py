# -*- coding: utf-8 -*-

import pandas as pd
import os
from wordcloud import WordCloud
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from PIL import ImageFont

def get_all_keyword_csv(directory):
	df_list = []

	csv_files = []
	for root, dirs, files in os.walk(directory):
		for file in files:
			if file.endswith(".csv"):
				csv_files.append(os.path.join(root, file))
	
	print("[ALL CSV FILE LIST]")
	print(csv_files)

	for idx, csv_file in enumerate(csv_files):
		result_df = pd.read_csv(csv_file, names=["Year", "keyword", "count"])
		result_df["keyword"] = result_df["keyword"].apply(lambda x: x.strip())
		result_df["keyword"] = result_df["keyword"].apply(lambda x: " ".join(x.replace("_", " ").split()))
		df_list.append(result_df)
	
	return df_list

def create_analysis_word_file(all_keyword_df_list, all_wordcloud_img_dir, yearly_publication_count_img_dir):
	# 0. Word File 생성
	doc = Document()

	# 1. 보고서 제목
	report_title = "연도별 논문 분석 결과"
	doc.add_heading(report_title, level=1)

	# 2. 연도별 논문 키워드 - 표 
	chapter1_title = "연도별 논문 키워드"
	doc.add_heading(chapter1_title, level=2)

	# 2.1 표 추가
	for keyword_df in all_keyword_df_list:
		# 2.1.1 표 생성
		doc.add_heading(f"{keyword_df['Year'][0]} Year Top 30 Keyword & Count", level=3)
		keyword_table = doc.add_table(rows=30, cols=2)
		keyword_table.style = doc.styles["Table Grid"]

		# 2.1.2 표 값 채우기
		for idx, row in keyword_df.iterrows():
			table_row = keyword_table.rows[idx].cells
			table_row[0].text = str(row["keyword"])
			table_row[1].text = str(row["count"])
	
	# 3. 연도별 논문 키워드 시각화 이미지 추가
	chapter2_title = "연도별 논문 키워드 시각화 - Word Cloud"
	doc.add_heading(chapter2_title, level=2)
	doc.add_picture(all_wordcloud_img_dir, width=Inches(6))

	# 4. 연도별 논문 게시수 그래프 이미지 추가
	chapter3_title = "연도별 논문 게시 수"
	doc.add_heading(chapter3_title, level=2)
	doc.add_picture(yearly_publication_count_img_dir, width=Inches(6))

	# 5. 보고서 저장
	doc.save("./Paper_Analysis_Result_Report.docx")

if __name__=="__main__":
	# 모든 결과 파일은 analysis_result 폴더에 .csv 파일로 있다고 가정

	# 1. 데이터 로드
	analysis_result_dir = "../analysis_result/abstract-keyword"
	all_keyword_df_list = get_all_keyword_csv(analysis_result_dir)

	# 2. pyplot 시각화 기본 세팅
	num_files = len(all_keyword_df_list) 
	num_cols = 3  # 한 줄에 보여줄 wordcloud 개수
	num_rows = (num_files + num_cols -1) // num_cols # all row 개수
	fig, axes = plt.subplots(num_rows, num_cols, figsize=(15, 5 * num_rows))

	# 3. keyword-csv 결과 word cloud
	for idx, keyword_df in enumerate(all_keyword_df_list):		
		# 3.1 word_freq dictionary 생성
		word_freq = {}
		for _, row in keyword_df.iterrows():
			keyword = row["keyword"]
			count = row["count"]
			word_freq[keyword] = count
		
		# 3.2 워드 클라우드 생성
		wc = WordCloud(width=500, height=500, max_words=200, background_color='white').generate_from_frequencies(word_freq)
		
		# 3.4 word cloud plot 추가
		row_pos = idx // num_cols
		col_pos = idx % num_cols

		axes[row_pos, col_pos].imshow(wc, interpolation='bilinear')
		axes[row_pos, col_pos].set_title(f"Word Cloud - {keyword_df['Year'][0]} Year Keyword")
		axes[row_pos, col_pos].axis('off')
		
	# 4. 시각화

	# 4.1 빈 공간 plot 축 제거
	for i in range(num_files, num_cols * num_rows):
		row_pos = i // num_cols
		col_pos = i % num_cols
		axes[row_pos, col_pos].axis('off')
	
	# 4.2 show plot
	plt.tight_layout()
	plt.savefig("./all_keyword_wordcloud.png")

	# 5. 연도별 논문 게시수 그래프 이미지 생성
	
	# 5.1 Load csv file
	count_file_dir = "../analysis_result/Yearly-Publication-Count/"
	file_list = os.listdir(count_file_dir)
	csv_file = [file for file in file_list if file.endswith(".csv")][0]
	csv_dir = os.path.join(count_file_dir, csv_file)
	print(csv_dir)
	paper_count_df = pd.read_csv(csv_dir, names=["Year", "count"])
	
	plt.figure(figsize=(10, 6))
	plt.plot(paper_count_df['Year'], paper_count_df['count'], marker='o', linestyle='-')
	plt.title('Yearly Publication Count')  # 그래프 제목
	plt.xlabel('Year')  # x축 레이블
	plt.ylabel('Count')  # y축 레이블
	plt.grid()  # 격자 표시
	plt.savefig("./yearly_publication_count.png")

	# 6. docx word 파일 생성
	all_wordcloud_img_dir = "./all_keyword_wordcloud.png"
	yearly_publication_count_img_dir = "./yearly_publication_count.png"
	create_analysis_word_file(all_keyword_df_list, all_wordcloud_img_dir, yearly_publication_count_img_dir)


